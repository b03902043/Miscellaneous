# -*- coding: utf-8 -*-
import docx
from copy import deepcopy

class WRow:
    def __init__(self, rw):
        self.row = rw
        self.tr = rw._tr

    def __setitem__(self, i, data):
        if isinstance(data, unicode):
            data = data.encode('utf-8')
        if not isinstance(data, str):
            data = str(data)
        data = data.decode('utf-8')
        if len(self.row.cells[i].paragraphs[0].runs) == 0:
            #print('[warn] add <run>')
            self.row.cells[i].paragraphs[0].add_run()
        self.row.cells[i].paragraphs[0].runs[0].text = data

    def __getitem__(self, i):
        return self.row.cells[i].paragraphs[0]

    def __len__(self):
        return len(self.row)

    def __str__(self):
        # merged cells will replicate cells ex. a | a | a | b
        #                                       ---- real ----
        #                                           a     | b
        return unicode(' | '.join(_.text for _ in self.row.cells)).encode('utf-8')

    def _add_col(self):
        _tc = deepcopy(self.tr.xpath('.//w:tc')[-1])
        self.tr.append(_tc)

class WTable:

    '''
    List of |WRow|
    '''

    def __init__(self, tbl):
        self._table = tbl
        self.table = map(WRow, tbl.rows)
        self.raw = tbl._tbl

    def __setitem__(self, pos, data):
        y, x = pos
        self.table[y][x] = data

    def __getitem__(self, pos):
        if isinstance(pos, tuple):
            y, x = pos
            return self.table[y][x]
        return self.table[pos]

    def __str__(self):
        return '\n'.join(str(r) for r in self.table)

    def __len__(self):
        return len(self.table)

    @property
    def obj(self):
        return deepcopy(self.raw)

    def remove(self):
        self.raw.getparent().remove(self.raw)

    def add_row(self, tr=None):
        _tr = deepcopy(tr if tr is not None else self.raw.xpath('.//w:tr')[-1])
        self.raw.append(_tr)
        self.table.append(WRow(docx.table._Row(_tr, self._table)))

    def add_col(self, tc=None):
        for r in self.table:
            r._add_col()
        self.raw.tblGrid.add_gridCol()
